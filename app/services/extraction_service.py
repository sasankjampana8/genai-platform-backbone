from dataclasses import dataclass
from io import BytesIO


@dataclass
class ExtractedPage:
    page_number: int
    text: str


@dataclass
class ExtractedDocument:
    pages: list[ExtractedPage]
    metadata: dict

    @property
    def full_text(self) -> str:
        return "\n\n".join(page.text for page in self.pages if page.text.strip())


class ExtractionService:
    """PDF/DOCX text extraction boundary."""

    def extract(self, file_bytes: bytes, file_name: str, content_type: str | None = None) -> ExtractedDocument:
        lower_name = file_name.lower()
        if lower_name.endswith(".pdf") or content_type == "application/pdf":
            return self._extract_pdf(file_bytes)
        if (
            lower_name.endswith(".docx")
            or content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            return self._extract_docx(file_bytes)
        raise ValueError(f"Unsupported document type for extraction: {file_name}")

    def _extract_pdf(self, file_bytes: bytes) -> ExtractedDocument:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError("pypdf is required for PDF extraction.") from exc

        reader = PdfReader(BytesIO(file_bytes))
        pages = [
            ExtractedPage(page_number=index, text=page.extract_text() or "")
            for index, page in enumerate(reader.pages, 1)
        ]
        return ExtractedDocument(pages=pages, metadata={"extractor": "pypdf", "page_count": len(pages)})

    def _extract_docx(self, file_bytes: bytes) -> ExtractedDocument:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("python-docx is required for DOCX extraction.") from exc

        document = Document(BytesIO(file_bytes))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        text = "\n".join(paragraphs)
        return ExtractedDocument(
            pages=[ExtractedPage(page_number=1, text=text)],
            metadata={"extractor": "python-docx", "page_count": 1, "paragraph_count": len(paragraphs)},
        )

